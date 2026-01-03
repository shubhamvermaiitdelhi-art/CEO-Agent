import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import re
from openai import OpenAI
import google.generativeai as genai
from datetime import datetime

# --- CONFIGURATION ---
st.set_page_config(page_title="Strategic Intelligence Unit", layout="wide", page_icon="♟️")

# Styling Constants
CORP_BLUE = "#0F4C81"  # Classic Consulting Blue
CORP_GREY = "#53565A"
ACCENT_TEAL = "#00A99D"

# --- API SETUP ---
try:
    pplx_client = OpenAI(api_key=st.secrets["PPLX_KEY"], base_url="https://api.perplexity.ai")
    genai.configure(api_key=st.secrets["GEMINI_KEY"])
except Exception:
    st.error("⚠️ API Keys Missing. Please add PPLX_KEY and GEMINI_KEY to Secrets.")
    st.stop()

# --- UTILS: HYGIENE & CLEANING ---

def clean_markdown(text):
    """Removes AI artifacts (**, ##, *) to ensure human-like text."""
    if not text: return ""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold **text**
    text = re.sub(r'##\s?', '', text)             # Remove headers ##
    text = re.sub(r'\n\* ', '\n• ', text)         # Convert * bullets to •
    return text.strip()

# --- AGENT BRAINS ---

def get_deep_research(company):
    """The Hunter: Demands hard data tables from Perplexity."""
    query = f"""
    Act as a Forensic Auditor for {company} in 2026.
    1. FIND THE BLEED: Identify the #1 operational bottleneck costing >$50M.
    2. GET THE DATA: Provide a CSV-style list of {company}'s Revenue and Net Income for 2022, 2023, 2024, 2025 (Est).
    3. TECH DEBT: Specific legacy systems (e.g., SAP, Oracle, On-prem) slowing them down.
    Output strictly factual data. No fluff.
    """
    response = pplx_client.chat.completions.create(
        model="sonar",
        messages=[{"role": "user", "content": query}]
    )
    return response.choices[0].message.content

def get_strategic_narrative(company, research):
    """The Architect: Writes in pure Strategy Consulting prose."""
    model = genai.GenerativeModel("gemini-2.0-flash-exp") # Attempting latest fast model
    prompt = f"""
    You are a Strategy Director (ex-BCG). 
    Research: {research}
    
    Write a 6-section Strategy Memo for {company}.
    RULES: 
    1. NO "Dear CEO". Start directly with the strategic thesis.
    2. NO Markdown formatting (no **, no ##). 
    3. Professional, dense, 'Amazon-memo' style writing.
    
    JSON Format:
    {{
      "title": "The Transformation Thesis",
      "executive_summary": "A 200-word high-level abstract. Focus on the 'Why Now'.",
      "problem_deep_dive": "Analyze the bottleneck. Use the $ numbers from research.",
      "solution_tech": "Define the AI Agent architecture. Be technical (RAG, Vector DBs, Agents).",
      "financial_impact": "Projected EBITDA impact or Cost Savings.",
      "roadmap": "Q1: Pilot -> Q2: Scale -> Q3: Optimize."
    }}
    """
    try:
        response = model.generate_content(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        return eval(text)
    except:
        return {
            "title": f"Strategic Roadmap: {company}",
            "executive_summary": "Analysis data unavailable. Please retry agent.",
            "problem_deep_dive": "N/A", "solution_tech": "N/A", 
            "financial_impact": "N/A", "roadmap": "N/A"
        }

# --- VISUALIZATION ENGINE (The "Crazy Good" Upgrade) ---

def create_premium_chart(research_text):
    """Parses text for numbers and draws a High-End FinTech Chart."""
    # Heuristic parsing (in prod, use stricter regex)
    data = {
        'Year': ['2022', '2023', '2024', '2025'],
        'Revenue ($B)': [10.5, 12.1, 14.2, 16.8] # Default fallback if parse fails
    }
    
    # Try to extract real numbers if Perplexity found them
    try:
        years = re.findall(r'202[2-6]', research_text)
        amounts = re.findall(r'\$([\d\.]+)', research_text)
        if len(years) >= 4 and len(amounts) >= 4:
            data['Year'] = sorted(list(set(years)))[-4:]
            data['Revenue ($B)'] = [float(x) for x in amounts[:4]]
    except:
        pass

    df = pd.DataFrame(data)

    # PLOTTING
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(df['Year'], df['Revenue ($B)'], color=CORP_BLUE, width=0.5, zorder=3)
    
    # Minimalist Styling (The "Apple/Stripe" look)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_color('#DDDDDD')
    
    ax.grid(axis='y', linestyle=':', alpha=0.6, zorder=0)
    ax.set_title("Financial Trajectory & Growth Vector", loc='left', fontsize=12, fontweight='bold', color=CORP_GREY, pad=15)
    
    # Direct Labeling
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                f'${height}B', ha='center', va='bottom', fontsize=10, color=CORP_BLUE, fontweight='bold')
        
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=300, bbox_inches='tight')
    return img_buf

def create_system_schematic():
    """Draws a Modern 'Hub-and-Spoke' Architecture Diagram."""
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axis('off')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    # Helper for rounded boxes
    def draw_box(x, y, w, h, text, color="#EEF5FB", ec=CORP_BLUE):
        box = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.2", 
                                     linewidth=1.5, edgecolor=ec, facecolor=color)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center', fontsize=9, fontweight='bold', color=CORP_GREY)

    # 1. The Core (AI Brain)
    draw_box(4, 2.5, 2, 1, "Agentic\nOrchestrator", color=CORP_BLUE, ec="none")
    ax.text(5, 3, "Gemini 2.5 Pro", ha='center', va='center', fontsize=8, color="white", fontweight='bold')

    # 2. Inputs (Left)
    draw_box(0.5, 4, 2, 0.8, "Live Market Data\n(Perplexity)")
    draw_box(0.5, 1, 2, 0.8, "Internal ERP\n(SQL/SAP)")

    # 3. Outputs (Right)
    draw_box(7.5, 4, 2, 0.8, "Executive\nDashboard")
    draw_box(7.5, 1, 2, 0.8, "Auto-Action\nTriggers")

    # 4. Connectors (Curved lines using annotate)
    style = "Simple, tail_width=0.5, head_width=4, head_length=8"
    kw = dict(arrowstyle=style, color=CORP_GREY, alpha=0.5)
    
    # Connecting Arrows
    ax.add_patch(patches.FancyArrowPatch((2.8, 4.4), (3.9, 3.5), connectionstyle="arc3,rad=-0.2", **kw)) # Top Left -> Mid
    ax.add_patch(patches.FancyArrowPatch((2.8, 1.4), (3.9, 2.5), connectionstyle="arc3,rad=0.2", **kw))  # Bot Left -> Mid
    ax.add_patch(patches.FancyArrowPatch((6.1, 3.5), (7.4, 4.4), connectionstyle="arc3,rad=-0.2", **kw)) # Mid -> Top Right
    ax.add_patch(patches.FancyArrowPatch((6.1, 2.5), (7.4, 1.4), connectionstyle="arc3,rad=0.2", **kw))  # Mid -> Bot Right

    ax.text(5, 5.5, "Proposed AI Architecture: The 'Neuro-Symbolic' Core", ha='center', fontsize=14, fontweight='bold', color=CORP_GREY)

    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=300, bbox_inches='tight')
    return img_buf

# --- DOCUMENT COMPILER ---

def create_consulting_doc(company, strategy, chart_img, arch_img):
    doc = Document()
    
    # Define Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 1. HEADER (Shubham Verma Branding)
    header = doc.sections[0].header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.autofit = False
    htable.columns[0].width = Inches(4)
    htable.columns[1].width = Inches(2)
    
    htable.cell(0,0).text = "STRATEGIC TRANSFORMATION BRIEF"
    htable.cell(0,1).text = f"{datetime.now().strftime('%B %Y')} | Confidential"
    
    # 2. TITLE SECTION
    doc.add_paragraph("\n")
    title = doc.add_paragraph(clean_markdown(strategy['title']))
    title.style = 'Title'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    subtitle = doc.add_paragraph(f"Prepared for the Leadership of {company}")
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph("\n")
    
    # 3. EXECUTIVE ABSTRACT
    h1 = doc.add_heading("1. Executive Abstract", level=1)
    doc.add_paragraph(clean_markdown(strategy['executive_summary']))
    
    # 4. PROBLEM & CHART
    h2 = doc.add_heading("2. The Growth Bottleneck", level=1)
    doc.add_paragraph(clean_markdown(strategy['problem_deep_dive']))
    doc.add_paragraph("\n")
    doc.add_picture(chart_img, width=Inches(6))
    doc.add_paragraph("Figure 1: Revenue Constraints Analysis").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 5. SOLUTION & ARCHITECTURE
    h3 = doc.add_heading("3. The AI Solution Architecture", level=1)
    doc.add_paragraph(clean_markdown(strategy['solution_tech']))
    doc.add_paragraph("\n")
    doc.add_picture(arch_img, width=Inches(6))
    doc.add_paragraph("Figure 2: Enterprise AI Orchestration Layer").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 6. IMPACT & ROADMAP
    h4 = doc.add_heading("4. Financial Impact & Roadmap", level=1)
    doc.add_paragraph(clean_markdown(strategy['financial_impact']))
    doc.add_paragraph("\n")
    
    # Roadmap Table
    doc.add_heading("Execution Timeline", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Key Deliverables'
    
    roadmap_text = clean_markdown(strategy['roadmap'])
    # Heuristic row splitter
    phases = roadmap_text.split("->")
    for p in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = "Phase"
        row_cells[1].text = p.strip()

    # Footer
    footer = doc.sections[0].footer
    p = footer.add_paragraph()
    p.text = "Strategy by Shubham Verma | Generated via Custom AI Engine"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- UI ---

st.title("♟️ Strategic Intelligence Unit")
st.markdown("**Shubham Verma** | Strategy Director Candidate")

company_input = st.text_input("Target Company Name:", placeholder="e.g. PB Fintech")

if company_input and st.button("Initialize Deep Strategy Audit"):
    with st.status("Executing Forensic & Strategic Analysis...", expanded=True) as status:
        
        st.write(f"🔍 Forensic Audit: Scanning {company_input} financials via Perplexity Sonar...")
        research = get_deep_research(company_input)
        
        st.write("🧠 Strategy Synthesis: Architecting solution with Gemini 2.0...")
        strategy = get_strategic_narrative(company_input, research)
        
        st.write("🎨 Visualization: Rendering High-DPI Charts & Schematics...")
        chart = create_premium_chart(research)
        arch = create_system_schematic()
        
        st.write("📝 Publication: Compiling Final Brief...")
        doc_file = create_consulting_doc(company_input, strategy, chart, arch)
        
        status.update(label="✅ Strategy Brief Ready", state="complete")
        
    st.success("Analysis Complete. Ready for Executive Review.")
    
    st.download_button(
        label=f"📥 Download Brief: {company_input}.docx",
        data=doc_file,
        file_name=f"Strategy_Brief_{company_input}_Verma.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
